import json
import re

from reportlab.lib import colors
from reportlab.lib.enums import (
    TA_CENTER,
    TA_LEFT,
)
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import (
    getSampleStyleSheet,
    ParagraphStyle,
)
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)

from docx import Document
from docx.shared import (
    Pt,
    Inches,
)
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
)

from app.services.resume_service import ResumeService


class ExportService:

    # =========================================================
    # DATA HELPERS
    # =========================================================

    @staticmethod
    def safe_json_load(value):

        try:
            return json.loads(value or "[]")

        except (
            json.JSONDecodeError,
            TypeError,
        ):
            return []

    @staticmethod
    def clean_text(text):

        if text is None:
            return ""

        return str(text).strip()

    @staticmethod
    def escape_html(text):

        text = str(text or "")

        return (
            text
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
        )

    @staticmethod
    def get_resume_data(resume_id):

        resume = ResumeService.get_resume(
            resume_id
        )

        if not resume:
            return None

        return {
            "title":
                resume.title,

            "template_name":
                resume.template_name
                or "Modern",

            "full_name":
                resume.full_name,

            "email":
                resume.email,

            "phone":
                resume.phone,

            "location":
                resume.location,

            "linkedin":
                resume.linkedin,

            "github":
                resume.github,

            "website":
                resume.website,

            "professional_summary":
                resume.professional_summary,

            "experience":
                ExportService.safe_json_load(
                    resume.experience
                ),

            "education":
                ExportService.safe_json_load(
                    resume.education
                ),

            "skills":
                resume.skills,

            "projects":
                ExportService.safe_json_load(
                    resume.projects
                ),

            "certifications":
                ExportService.safe_json_load(
                    resume.certifications
                ),

            "languages":
                resume.languages,
        }

    @staticmethod
    def safe_filename(name):

        name = name or "resume"

        name = re.sub(
            r'[\\/*?:"<>|]',
            "",
            name
        )

        name = name.strip()

        return name or "resume"

    @staticmethod
    def get_template_name(data):

        return (
            data.get(
                "template_name",
                "Modern"
            )
            .strip()
            .lower()
        )

    @staticmethod
    def get_contact_items(data):

        contact_items = []

        fields = [
            "email",
            "phone",
            "location",
            "linkedin",
            "github",
            "website",
        ]

        for field in fields:

            value = ExportService.clean_text(
                data.get(field, "")
            )

            if value:
                contact_items.append(
                    value
                )

        return contact_items

    # =========================================================
    # PDF STYLE FACTORY
    # =========================================================

    @staticmethod
    def get_pdf_styles(template_name):

        styles = getSampleStyleSheet()

        if template_name == "classic":

            return {
                "name": ParagraphStyle(
                    "ClassicName",
                    parent=styles["Heading1"],
                    fontName="Times-Bold",
                    fontSize=22,
                    leading=27,
                    alignment=TA_CENTER,
                    spaceAfter=5,
                ),

                "contact": ParagraphStyle(
                    "ClassicContact",
                    parent=styles["Normal"],
                    fontName="Times-Roman",
                    fontSize=9,
                    leading=12,
                    alignment=TA_CENTER,
                    spaceAfter=12,
                ),

                "section": ParagraphStyle(
                    "ClassicSection",
                    parent=styles["Heading2"],
                    fontName="Times-Bold",
                    fontSize=12,
                    leading=15,
                    spaceBefore=12,
                    spaceAfter=5,
                    textColor=colors.black,
                ),

                "normal": ParagraphStyle(
                    "ClassicNormal",
                    parent=styles["Normal"],
                    fontName="Times-Roman",
                    fontSize=9.5,
                    leading=14,
                    spaceAfter=5,
                ),

                "item_title": ParagraphStyle(
                    "ClassicItem",
                    parent=styles["Normal"],
                    fontName="Times-Bold",
                    fontSize=10,
                    leading=13,
                    spaceBefore=4,
                    spaceAfter=2,
                ),

                "accent":
                    colors.black,
            }

        if template_name == "minimal":

            return {
                "name": ParagraphStyle(
                    "MinimalName",
                    parent=styles["Heading1"],
                    fontSize=21,
                    leading=25,
                    alignment=TA_LEFT,
                    spaceAfter=4,
                    textColor=colors.HexColor(
                        "#222222"
                    ),
                ),

                "contact": ParagraphStyle(
                    "MinimalContact",
                    parent=styles["Normal"],
                    fontSize=8.5,
                    leading=11,
                    alignment=TA_LEFT,
                    spaceAfter=12,
                    textColor=colors.HexColor(
                        "#666666"
                    ),
                ),

                "section": ParagraphStyle(
                    "MinimalSection",
                    parent=styles["Heading2"],
                    fontSize=11,
                    leading=14,
                    spaceBefore=10,
                    spaceAfter=5,
                    textColor=colors.HexColor(
                        "#555555"
                    ),
                ),

                "normal": ParagraphStyle(
                    "MinimalNormal",
                    parent=styles["Normal"],
                    fontSize=9,
                    leading=13,
                    spaceAfter=4,
                ),

                "item_title": ParagraphStyle(
                    "MinimalItem",
                    parent=styles["Normal"],
                    fontSize=9.5,
                    leading=12,
                    spaceBefore=4,
                    spaceAfter=2,
                ),

                "accent":
                    colors.HexColor(
                        "#777777"
                    ),
            }

        return {
            "name": ParagraphStyle(
                "ModernName",
                parent=styles["Heading1"],
                fontSize=23,
                leading=28,
                alignment=TA_CENTER,
                spaceAfter=6,
                textColor=colors.HexColor(
                    "#1F2937"
                ),
            ),

            "contact": ParagraphStyle(
                "ModernContact",
                parent=styles["Normal"],
                fontSize=9,
                leading=12,
                alignment=TA_CENTER,
                spaceAfter=14,
                textColor=colors.HexColor(
                    "#4B5563"
                ),
            ),

            "section": ParagraphStyle(
                "ModernSection",
                parent=styles["Heading2"],
                fontSize=12,
                leading=15,
                spaceBefore=11,
                spaceAfter=6,
                textColor=colors.HexColor(
                    "#2563EB"
                ),
            ),

            "normal": ParagraphStyle(
                "ModernNormal",
                parent=styles["Normal"],
                fontSize=9.5,
                leading=14,
                spaceAfter=5,
            ),

            "item_title": ParagraphStyle(
                "ModernItem",
                parent=styles["Normal"],
                fontSize=10,
                leading=13,
                spaceBefore=4,
                spaceAfter=2,
            ),

            "accent":
                colors.HexColor(
                    "#2563EB"
                ),
        }

    # =========================================================
    # PDF EXPORT
    # =========================================================

    @staticmethod
    def export_pdf(
        resume_id,
        file_path
    ):

        data = (
            ExportService.get_resume_data(
                resume_id
            )
        )

        if not data:

            raise ValueError(
                "Resume not found."
            )

        template_name = (
            ExportService.get_template_name(
                data
            )
        )

        document = SimpleDocTemplate(
            file_path,
            pagesize=A4,
            rightMargin=45,
            leftMargin=45,
            topMargin=40,
            bottomMargin=40,
        )

        styles = (
            ExportService.get_pdf_styles(
                template_name
            )
        )

        story = []

        full_name = (
            ExportService.clean_text(
                data["full_name"]
            )
            or "YOUR NAME"
        )

        # -----------------------------------------------------
        # MODERN HEADER
        # -----------------------------------------------------

        if template_name == "modern":

            header_table = Table(
                [[
                    Paragraph(
                        ExportService.escape_html(
                            full_name
                        ),
                        styles["name"]
                    )
                ]],
                colWidths=[6.5 * inch],
            )

            header_table.setStyle(
                TableStyle([
                    (
                        "BACKGROUND",
                        (0, 0),
                        (-1, -1),
                        colors.HexColor(
                            "#EFF6FF"
                        ),
                    ),
                    (
                        "BOX",
                        (0, 0),
                        (-1, -1),
                        0.8,
                        styles["accent"],
                    ),
                    (
                        "TOPPADDING",
                        (0, 0),
                        (-1, -1),
                        12,
                    ),
                    (
                        "BOTTOMPADDING",
                        (0, 0),
                        (-1, -1),
                        8,
                    ),
                ])
            )

            story.append(
                header_table
            )

        else:

            story.append(
                Paragraph(
                    ExportService.escape_html(
                        full_name
                    ),
                    styles["name"]
                )
            )

        contact_items = (
            ExportService.get_contact_items(
                data
            )
        )

        if contact_items:

            story.append(
                Paragraph(
                    ExportService.escape_html(
                        " • ".join(
                            contact_items
                        )
                    ),
                    styles["contact"]
                )
            )

        # -----------------------------------------------------
        # CONTENT
        # -----------------------------------------------------

        summary = (
            ExportService.clean_text(
                data[
                    "professional_summary"
                ]
            )
        )

        if summary:

            ExportService.add_pdf_section(
                story,
                ExportService.get_section_title(
                    "summary",
                    template_name
                ),
                styles["section"],
                template_name
            )

            story.append(
                Paragraph(
                    ExportService.escape_html(
                        summary
                    ).replace(
                        "\n",
                        "<br/>"
                    ),
                    styles["normal"]
                )
            )

        ExportService.add_pdf_entry_group(
            story,
            "experience",
            data["experience"],
            styles,
            template_name
        )

        ExportService.add_pdf_entry_group(
            story,
            "education",
            data["education"],
            styles,
            template_name
        )

        skills = (
            ExportService.clean_text(
                data["skills"]
            )
        )

        if skills:

            ExportService.add_pdf_section(
                story,
                ExportService.get_section_title(
                    "skills",
                    template_name
                ),
                styles["section"],
                template_name
            )

            story.append(
                Paragraph(
                    ExportService.escape_html(
                        skills
                    ).replace(
                        "\n",
                        "<br/>"
                    ),
                    styles["normal"]
                )
            )

        ExportService.add_pdf_entry_group(
            story,
            "projects",
            data["projects"],
            styles,
            template_name
        )

        ExportService.add_pdf_entry_group(
            story,
            "certifications",
            data["certifications"],
            styles,
            template_name
        )

        languages = (
            ExportService.clean_text(
                data["languages"]
            )
        )

        if languages:

            ExportService.add_pdf_section(
                story,
                ExportService.get_section_title(
                    "languages",
                    template_name
                ),
                styles["section"],
                template_name
            )

            story.append(
                Paragraph(
                    ExportService.escape_html(
                        languages
                    ).replace(
                        "\n",
                        "<br/>"
                    ),
                    styles["normal"]
                )
            )

        document.build(
            story
        )

        return file_path

    # =========================================================
    # PDF HELPERS
    # =========================================================

    @staticmethod
    def get_section_title(
        section,
        template_name
    ):

        modern_titles = {
            "summary":
                "PROFESSIONAL SUMMARY",

            "experience":
                "EXPERIENCE",

            "education":
                "EDUCATION",

            "skills":
                "SKILLS",

            "projects":
                "PROJECTS",

            "certifications":
                "CERTIFICATIONS",

            "languages":
                "LANGUAGES",
        }

        minimal_titles = {
            "summary":
                "Summary",

            "experience":
                "Experience",

            "education":
                "Education",

            "skills":
                "Skills",

            "projects":
                "Projects",

            "certifications":
                "Certifications",

            "languages":
                "Languages",
        }

        if template_name == "minimal":
            return minimal_titles[section]

        return modern_titles[section]

    @staticmethod
    def add_pdf_section(
        story,
        title,
        style,
        template_name
    ):

        story.append(
            Spacer(
                1,
                5
            )
        )

        story.append(
            Paragraph(
                title,
                style
            )
        )

        if template_name == "classic":

            story.append(
                Table(
                    [[""]],
                    colWidths=[
                        6.45 * inch
                    ],
                    rowHeights=[1],
                    style=TableStyle([
                        (
                            "BACKGROUND",
                            (0, 0),
                            (-1, -1),
                            colors.black,
                        )
                    ])
                )
            )

    @staticmethod
    def add_pdf_entry_group(
        story,
        section,
        entries,
        styles,
        template_name
    ):

        if not entries:
            return

        ExportService.add_pdf_section(
            story,
            ExportService.get_section_title(
                section,
                template_name
            ),
            styles["section"],
            template_name
        )

        for entry in entries:

            ExportService.add_pdf_entry(
                story,
                entry,
                styles["normal"],
                styles["item_title"]
            )

    @staticmethod
    def add_pdf_entry(
        story,
        entry,
        normal_style,
        title_style
    ):

        title = (
            ExportService.clean_text(
                entry.get(
                    "title",
                    ""
                )
            )
        )

        subtitle = (
            ExportService.clean_text(
                entry.get(
                    "subtitle",
                    ""
                )
            )
        )

        date = (
            ExportService.clean_text(
                entry.get(
                    "date",
                    ""
                )
            )
        )

        description = (
            ExportService.clean_text(
                entry.get(
                    "description",
                    ""
                )
            )
        )

        heading = (
            ExportService.escape_html(
                title
            )
        )

        if subtitle:

            heading += (
                " — "
                + ExportService.escape_html(
                    subtitle
                )
            )

        if date:

            heading += (
                " | "
                + ExportService.escape_html(
                    date
                )
            )

        if heading:

            story.append(
                Paragraph(
                    f"<b>{heading}</b>",
                    title_style
                )
            )

        if description:

            story.append(
                Paragraph(
                    ExportService.escape_html(
                        description
                    ).replace(
                        "\n",
                        "<br/>"
                    ),
                    normal_style
                )
            )

    # =========================================================
    # DOCX EXPORT
    # =========================================================

    @staticmethod
    def export_docx(
        resume_id,
        file_path
    ):

        data = (
            ExportService.get_resume_data(
                resume_id
            )
        )

        if not data:

            raise ValueError(
                "Resume not found."
            )

        template_name = (
            ExportService.get_template_name(
                data
            )
        )

        document = Document()

        section = (
            document.sections[0]
        )

        section.top_margin = Inches(
            0.55
        )

        section.bottom_margin = Inches(
            0.55
        )

        section.left_margin = Inches(
            0.65
        )

        section.right_margin = Inches(
            0.65
        )

        # -----------------------------------------------------
        # NAME
        # -----------------------------------------------------

        name = (
            ExportService.clean_text(
                data["full_name"]
            )
            or "YOUR NAME"
        )

        paragraph = (
            document.add_paragraph()
        )

        if template_name == "minimal":

            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT
            )

            name_size = 20

        else:

            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
            )

            name_size = 22

        run = paragraph.add_run(
            name
        )

        run.bold = True
        run.font.size = Pt(
            name_size
        )

        if template_name == "classic":

            run.font.name = (
                "Times New Roman"
            )

        # -----------------------------------------------------
        # CONTACT
        # -----------------------------------------------------

        contact_items = (
            ExportService.get_contact_items(
                data
            )
        )

        if contact_items:

            paragraph = (
                document.add_paragraph()
            )

            if template_name == "minimal":

                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.LEFT
                )

            else:

                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER
                )

            run = paragraph.add_run(
                " • ".join(
                    contact_items
                )
            )

            run.font.size = Pt(9)

        # -----------------------------------------------------
        # SUMMARY
        # -----------------------------------------------------

        summary = (
            ExportService.clean_text(
                data[
                    "professional_summary"
                ]
            )
        )

        if summary:

            ExportService.add_docx_heading(
                document,
                ExportService.get_section_title(
                    "summary",
                    template_name
                ),
                template_name
            )

            document.add_paragraph(
                summary
            )

        # -----------------------------------------------------
        # EXPERIENCE
        # -----------------------------------------------------

        ExportService.add_docx_entry_group(
            document,
            "experience",
            data["experience"],
            template_name
        )

        # -----------------------------------------------------
        # EDUCATION
        # -----------------------------------------------------

        ExportService.add_docx_entry_group(
            document,
            "education",
            data["education"],
            template_name
        )

        # -----------------------------------------------------
        # SKILLS
        # -----------------------------------------------------

        skills = (
            ExportService.clean_text(
                data["skills"]
            )
        )

        if skills:

            ExportService.add_docx_heading(
                document,
                ExportService.get_section_title(
                    "skills",
                    template_name
                ),
                template_name
            )

            document.add_paragraph(
                skills
            )

        # -----------------------------------------------------
        # PROJECTS
        # -----------------------------------------------------

        ExportService.add_docx_entry_group(
            document,
            "projects",
            data["projects"],
            template_name
        )

        # -----------------------------------------------------
        # CERTIFICATIONS
        # -----------------------------------------------------

        ExportService.add_docx_entry_group(
            document,
            "certifications",
            data["certifications"],
            template_name
        )

        # -----------------------------------------------------
        # LANGUAGES
        # -----------------------------------------------------

        languages = (
            ExportService.clean_text(
                data["languages"]
            )
        )

        if languages:

            ExportService.add_docx_heading(
                document,
                ExportService.get_section_title(
                    "languages",
                    template_name
                ),
                template_name
            )

            document.add_paragraph(
                languages
            )

        document.save(
            file_path
        )

        return file_path

    # =========================================================
    # DOCX HELPERS
    # =========================================================

    @staticmethod
    def add_docx_heading(
        document,
        title,
        template_name
    ):

        paragraph = (
            document.add_paragraph()
        )

        run = paragraph.add_run(
            title
        )

        run.bold = True

        if template_name == "minimal":

            run.font.size = Pt(11)

        else:

            run.font.size = Pt(12)

        if template_name == "classic":

            run.font.name = (
                "Times New Roman"
            )

        paragraph.paragraph_format.space_before = (
            Pt(10)
        )

        paragraph.paragraph_format.space_after = (
            Pt(4)
        )

        return paragraph

    @staticmethod
    def add_docx_entry_group(
        document,
        section,
        entries,
        template_name
    ):

        if not entries:
            return

        ExportService.add_docx_heading(
            document,
            ExportService.get_section_title(
                section,
                template_name
            ),
            template_name
        )

        for entry in entries:

            ExportService.add_docx_entry(
                document,
                entry,
                template_name
            )

    @staticmethod
    def add_docx_entry(
        document,
        entry,
        template_name
    ):

        title = (
            ExportService.clean_text(
                entry.get(
                    "title",
                    ""
                )
            )
        )

        subtitle = (
            ExportService.clean_text(
                entry.get(
                    "subtitle",
                    ""
                )
            )
        )

        date = (
            ExportService.clean_text(
                entry.get(
                    "date",
                    ""
                )
            )
        )

        description = (
            ExportService.clean_text(
                entry.get(
                    "description",
                    ""
                )
            )
        )

        paragraph = (
            document.add_paragraph()
        )

        run = paragraph.add_run(
            title
        )

        run.bold = True
        run.font.size = Pt(10)

        if template_name == "classic":

            run.font.name = (
                "Times New Roman"
            )

        details = []

        if subtitle:
            details.append(
                subtitle
            )

        if date:
            details.append(
                date
            )

        if details:

            details_paragraph = (
                document.add_paragraph()
            )

            details_run = (
                details_paragraph.add_run(
                    " | ".join(
                        details
                    )
                )
            )

            details_run.italic = True
            details_run.font.size = Pt(9)

        if description:

            description_paragraph = (
                document.add_paragraph(
                    description
                )
            )

            description_paragraph.paragraph_format.space_after = (
                Pt(4)
            )